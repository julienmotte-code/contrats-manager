"""
Service de génération de contrats Word par publipostage.
Remplace les champs «ChampNom» dans les modèles DOCX par les valeurs réelles.
"""
import re
import logging
from datetime import date
from pathlib import Path
from typing import Optional, Dict

from docx import Document
from sqlalchemy.orm import Session

from app.models.models import Contrat, ClientCache, DocumentGenere, ModeleDocument

logger = logging.getLogger(__name__)

STORAGE_DIR   = Path("/app/storage")
MODELES_DIR   = STORAGE_DIR / "modeles"
DOCUMENTS_DIR = STORAGE_DIR / "documents_generes"

FAMILLE_MODELE = {
    "COSOLUCE":       "Modele_Contrat_Cosoluce_et_Annexes.docx",
    "CANTINE":        "Modele_Contrat_Cantine_de_France.docx",
    "MAINTENANCE":    "Modele_Contrat_Maintenance_Systeme.docx",
    "ASSISTANCE_TEL": "Modele_Contrat_Assistance_Cityweb.docx",
}

FAMILLE_LABEL = {
    "COSOLUCE":       "Abonnement/Maintenance Cosoluce",
    "CANTINE":        "Abonnement Cantine de France",
    "MAINTENANCE":    "Maintenance Système",
    "ASSISTANCE_TEL": "Assistance Téléphonique",
    "DIGITECH":       "Digitech",
    "KIWI_BACKUP":    "Kiwi Backup",
}

CHAMPS = {
    "NomClient":         ["NomClient", "NomSite"],
    "AdresseClient":     ["AdresseClient", "AdrSite"],
    "CPClient":          ["CPClient", "CPSite"],
    "VilleClient":       ["VilleClient", "VilleSite"],
    "SIRETClient":       ["SIRETClient", "CodeSIRET"],
    "EmailClient":       ["EmailClient", "Email"],
    "TelClient":         ["TelClient"],
    "IdSite":            ["IdSite", "RefClient"],
    "NoContrat":         ["NoContrat"],
    "CodeInstance":      ["CodeInstance"],
    "LibCodePrestation": ["LibCodePrestation"],
    "NoAbonnement":      ["NoAbonnement", "NoContratCosoluce"],
    "DateDebut":         ["DateDebut", "DateContrat", "DateDebutContrat"],
    "DateFin":           ["DateFin", "DateFinContrat"],
    "DateDoc":           ["DateDoc", "DateDuJour", "DatduJour"],
    "DateSignature":     ["DateSignature"],
    "DatePremFactu":     ["DatePremFactu"],
    "MontantHT":         ["MontantHT", "PrixAnnuelContrat"],
    "TotalHT":           ["TotalHT", "TotalAnnuelContrat"],
    "NomSignataire":     ["NomSignataire"],
    "QualiteSignataire": ["QualiteSignataire"],
    "NomInterlocuteur":  ["NomInterlocuteur"],
    "Progiciel1":        ["Progiciel1"],
    "DescProg1":         ["DescriptionProgiciel1"],
    "PrixProg1":         ["PrixProgiciel1"],
    "Progiciel2":        ["Progiciel2"],
    "DescProg2":         ["DescriptionProgiciel2"],
    "PrixProg2":         ["PrixProgiciel2"],
}

_RE_COL  = re.compile(r"\xabCOL\d+(?:IdSite|NomSite)\xbb")
_RE_REST = re.compile(r"\xab[^\xbb]+\xbb")


def _fmt_date(d) -> str:
    if not d:
        return ""
    if isinstance(d, str):
        try:
            d = date.fromisoformat(d)
        except Exception:
            return d
    return d.strftime("%d/%m/%Y")


def _fmt_montant(v) -> str:
    if v is None:
        return ""
    try:
        return f"{float(v):,.2f}".replace(",", "\u00a0").replace(".", ",")
    except Exception:
        return str(v)


def _construire_variables(contrat: Contrat, client: Optional[ClientCache]) -> Dict[str, str]:
    articles = sorted(contrat.articles or [], key=lambda a: a.rang)
    art0 = next((a for a in articles if a.rang == 0), None)
    art1 = next((a for a in articles if a.rang == 1), None)
    montant = float(contrat.montant_annuel_ht) if contrat.montant_annuel_ht else 0.0
    today   = _fmt_date(date.today())
    return {
        "NomClient":         (client.nom if client else contrat.client_nom) or "",
        "AdresseClient":     (client.adresse_ligne1 if client else "") or "",
        "CPClient":          (client.code_postal if client else "") or "",
        "VilleClient":       (client.ville if client else "") or "",
        "SIRETClient":       (client.siret if client else "") or "",
        "EmailClient":       (client.email if client else "") or "",
        "TelClient":         (client.telephone if client else "") or "",
        "IdSite":            (client.numero_client if client else "") or "",
        "NoContrat":         contrat.numero_contrat or "",
        "CodeInstance":      contrat.numero_contrat or "",
        "LibCodePrestation": FAMILLE_LABEL.get(contrat.famille_contrat, contrat.famille_contrat or ""),
        "NoAbonnement":      f"AB-{contrat.numero_contrat}",
        "DateDebut":         _fmt_date(contrat.date_debut),
        "DateFin":           _fmt_date(contrat.date_fin),
        "DateDoc":           today,
        "DateSignature":     today,
        "DatePremFactu":     _fmt_date(contrat.date_debut),
        "MontantHT":         _fmt_montant(montant),
        "TotalHT":           _fmt_montant(montant),
        "NomSignataire":     (client.contact_nom if client else "") or "",
        "QualiteSignataire": (client.contact_fonction if client else "") or "",
        "NomInterlocuteur":  (client.contact_nom if client else "") or "",
        "Progiciel1":        art0.designation if art0 else "",
        "DescProg1":         art0.designation if art0 else "",
        "PrixProg1":         _fmt_montant(float(art0.prix_unitaire_ht) * float(art0.quantite or 1)) if art0 and art0.prix_unitaire_ht else "",
        "Progiciel2":        art1.designation if art1 else "",
        "DescProg2":         art1.designation if art1 else "",
        "PrixProg2":         _fmt_montant(float(art1.prix_unitaire_ht) * float(art1.quantite or 1)) if art1 and art1.prix_unitaire_ht else "",
    }


def _remplacer_texte(texte: str, variables: Dict[str, str]) -> str:
    for canon, aliases in CHAMPS.items():
        valeur = variables.get(canon, "")
        for alias in aliases:
            texte = texte.replace(f"\xab{alias}\xbb", valeur)
    texte = _RE_COL.sub("", texte)
    texte = _RE_REST.sub("", texte)
    return texte


def _traiter_paragraphe(para, variables: Dict[str, str]):
    texte_complet = "".join(r.text for r in para.runs)
    if "\xab" not in texte_complet:
        return
    for run in para.runs:
        if "\xab" in run.text:
            run.text = _remplacer_texte(run.text, variables)
            run.font.color.rgb = None
    texte_apres = "".join(r.text for r in para.runs)
    if "\xab" not in texte_apres:
        return
    texte_reconstruit = _remplacer_texte(texte_apres, variables)
    if para.runs:
        para.runs[0].text = texte_reconstruit
        para.runs[0].font.color.rgb = None
        for run in para.runs[1:]:
            run.text = ""


def _traiter_tableau(table, variables: Dict[str, str]):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _traiter_paragraphe(para, variables)
            for sous_tableau in cell.tables:
                _traiter_tableau(sous_tableau, variables)


def _traiter_document(doc: Document, variables: Dict[str, str]):
    for para in doc.paragraphs:
        _traiter_paragraphe(para, variables)
    for table in doc.tables:
        _traiter_tableau(table, variables)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _traiter_paragraphe(para, variables)
        for table in section.header.tables:
            _traiter_tableau(table, variables)


def _trouver_modele(famille: str, db: Session) -> Optional[Path]:
    type_doc  = f"CONTRAT_{famille}"
    modele_db = (
        db.query(ModeleDocument)
        .filter(ModeleDocument.type_document == type_doc, ModeleDocument.actif == True)
        .order_by(ModeleDocument.uploaded_at.desc())
        .first()
    )
    if modele_db:
        p = Path(modele_db.chemin_fichier)
        if p.exists():
            return p
    nom = FAMILLE_MODELE.get(famille)
    if nom:
        p = MODELES_DIR / nom
        if p.exists():
            return p
    return None


def generer_document(contrat: Contrat, client: Optional[ClientCache], db: Session, generated_by: str = "system") -> Dict:
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    modele_path = _trouver_modele(contrat.famille_contrat, db)
    if not modele_path:
        return {"success": False, "error": f"Aucun modèle disponible pour '{contrat.famille_contrat}'. Déposez un .docx dans /app/storage/modeles/."}
    variables = _construire_variables(contrat, client)
    try:
        doc = Document(str(modele_path))
    except Exception as e:
        return {"success": False, "error": f"Impossible d'ouvrir le modèle : {e}"}
    _traiter_document(doc, variables)
    nom_client_safe = ((client.nom if client else contrat.client_nom) or "client").replace(" ", "_")[:30]
    nom_fichier = f"Contrat_{contrat.numero_contrat.replace('/', '-')}_{nom_client_safe}_{date.today().strftime('%Y%m%d')}.docx"
    chemin = DOCUMENTS_DIR / nom_fichier
    try:
        doc.save(str(chemin))
    except Exception as e:
        return {"success": False, "error": f"Impossible de sauvegarder : {e}"}
    doc_genere = DocumentGenere(
        contrat_id=contrat.id,
        type_document=f"CONTRAT_{contrat.famille_contrat}",
        nom_fichier=nom_fichier,
        chemin_docx=str(chemin),
        modele_utilise=str(modele_path),
        variables_json=variables,
        generated_by=generated_by,
    )
    db.add(doc_genere)
    db.commit()
    db.refresh(doc_genere)
    logger.info(f"Document généré : {nom_fichier} (contrat {contrat.numero_contrat})")
    return {"success": True, "document_id": str(doc_genere.id), "nom_fichier": nom_fichier}


def lister_documents_contrat(contrat_id: str, db: Session) -> list:
    docs = (
        db.query(DocumentGenere)
        .filter(DocumentGenere.contrat_id == contrat_id)
        .order_by(DocumentGenere.generated_at.desc())
        .all()
    )
    return [
        {
            "id": str(d.id),
            "type_document": d.type_document,
            "nom_fichier": d.nom_fichier,
            "generated_by": d.generated_by,
            "generated_at": d.generated_at.isoformat() if d.generated_at else None,
        }
        for d in docs
    ]
