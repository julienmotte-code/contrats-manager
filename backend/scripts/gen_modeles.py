from docx import Document
from docx.shared import Pt, RGBColor, Cm
from pathlib import Path

RED  = RGBColor(0xCC, 0x00, 0x00)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
OUT  = Path("/app/storage/modeles")
OUT.mkdir(parents=True, exist_ok=True)

def champ(para, nom):
    r = para.add_run(f"\u00ab{nom}\u00bb")
    r.font.color.rgb = RED; r.font.bold = True; r.font.name = "Arial"; r.font.size = Pt(10)
    return r

def txt(para, texte, bold=False, color=None, size=10):
    r = para.add_run(texte)
    r.font.name = "Arial"; r.font.size = Pt(size); r.font.bold = bold
    if color: r.font.color.rgb = color
    return r

def titre(doc, texte):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(texte)
    r.font.name = "Arial"; r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = BLUE
    return p

def entete_tableau(table, *cols):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    row = table.rows[0]
    for i, col in enumerate(cols):
        cell = row.cells[i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(col)
        r.font.bold = True; r.font.name = "Arial"; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F4E79'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

def bloc_signataires(doc):
    p = doc.add_paragraph()
    txt(p, "Fait \u00e0 Villeneuve d\u2019Ascq, le "); champ(p, "DateDoc"); txt(p, " \u2014 En 2 exemplaires.")
    t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
    c0, c1 = t.rows[0].cells
    p2 = c0.paragraphs[0]; txt(p2, "Le Client", bold=True)
    p3 = c0.add_paragraph(); txt(p3, "Signature et cachet \u2014 \u00ab lu et approuv\u00e9 \u00bb")
    p4 = c0.add_paragraph(); champ(p4, "NomSignataire"); txt(p4, " \u2014 "); champ(p4, "QualiteSignataire")
    p5 = c1.paragraphs[0]; txt(p5, "Le Prestataire", bold=True)
    c1.add_paragraph().add_run("SGI Informatique \u2014 G\u00e9rant").font.name = "Arial"

def entete_client(doc, champs_supp=None):
    t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
    c0, c1 = t.rows[0].cells
    txt(c0.paragraphs[0], "LE PRESTATAIRE", bold=True, color=BLUE)
    for line in ["SARL S.G.I", "9 Avenue de la Cr\u00e9ativit\u00e9", "59650 VILLENEUVE D\u2019ASCQ", "SIRET 531 891 307 00021"]:
        c0.add_paragraph().add_run(line).font.name = "Arial"
    txt(c1.paragraphs[0], "LE CLIENT", bold=True, color=BLUE)
    for cn in ["NomClient", "AdresseClient", "CPClient", "VilleClient"]:
        p2 = c1.add_paragraph(); champ(p2, cn)
    p2 = c1.add_paragraph(); txt(p2, "SIRET : "); champ(p2, "SIRETClient")
    if champs_supp:
        for label, nom_champ in champs_supp:
            p2 = c1.add_paragraph(); txt(p2, label); champ(p2, nom_champ)
    return t

# ═══════════════════════════════════════════════════
# COSOLUCE
# ═══════════════════════════════════════════════════
doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2); s.bottom_margin=Cm(2); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
p = doc.add_paragraph(); txt(p, "SGI INFORMATIQUE \u2014 Contrat ", bold=True, color=BLUE, size=13); champ(p, "LibCodePrestation"); txt(p, " \u2014 ", color=BLUE, size=13); champ(p, "NoContrat")
doc.add_paragraph()
entete_client(doc, [("T\u00e9l : ", "TelClient"), ("Email : ", "EmailClient")])
doc.add_paragraph()
p = doc.add_paragraph(); txt(p, "R\u00e9f\u00e9rence : "); champ(p, "NoContrat"); txt(p, " \u2014 Date : "); champ(p, "DateDoc")
p = doc.add_paragraph(); txt(p, "Code instance : "); champ(p, "CodeInstance")
titre(doc, "ARTICLE 1 \u2014 OBJET")
p = doc.add_paragraph(); txt(p, "Abonnement/maintenance des progiciels Cosoluce. Progiciel principal : "); champ(p, "Progiciel1")
titre(doc, "ARTICLE 2 \u2014 DUR\u00c9E")
p = doc.add_paragraph(); txt(p, "Du "); champ(p, "DateDebut"); txt(p, " au "); champ(p, "DateFin"); txt(p, ". Reconduction annuelle tacite sauf d\u00e9nonciation 3 mois avant \u00e9ch\u00e9ance.")
titre(doc, "ARTICLE 3 \u2014 CONDITIONS FINANCI\u00c8RES")
t2 = doc.add_table(rows=4, cols=2); t2.style = "Table Grid"
entete_tableau(t2, "D\u00e9signation", "Montant HT/an")
for i, (lab, cp) in enumerate([("", "Progiciel1"), ("", "Progiciel2")], 1):
    p0 = t2.rows[i].cells[0].paragraphs[0]; p0.clear(); champ(p0, cp)
    p1 = t2.rows[i].cells[1].paragraphs[0]; p1.clear()
    champ(p1, "PrixProgiciel" + str(i)); txt(p1, " \u20ac")
p0 = t2.rows[3].cells[0].paragraphs[0]; p0.clear(); txt(p0, "TOTAL HT", bold=True)
p1 = t2.rows[3].cells[1].paragraphs[0]; p1.clear(); champ(p1, "MontantHT"); txt(p1, " \u20ac")
p = doc.add_paragraph(); txt(p, "Premi\u00e8re facturation le "); champ(p, "DatePremFactu"); txt(p, ".")
titre(doc, "ARTICLE 4 \u2014 R\u00c9VISION")
doc.add_paragraph().add_run("R\u00e9vision annuelle Syntec Ao\u00fbt : Prix N = Prix N-1 \u00d7 (Ao\u00fbt N-1 / Ao\u00fbt N-2).").font.name = "Arial"
titre(doc, "ARTICLE 5 \u2014 PRESTATIONS")
doc.add_paragraph().add_run("MAJ r\u00e9glementaires, assistance t\u00e9l\u00e9phonique heures ouvrables, hotline N1/N2, h\u00e9bergement s\u00e9curis\u00e9 France.").font.name = "Arial"
titre(doc, "ARTICLE 6 \u2014 R\u00c9SILIATION")
doc.add_paragraph().add_run("Pr\u00e9avis 3 mois par lettre recommand\u00e9e avant \u00e9ch\u00e9ance annuelle.").font.name = "Arial"
titre(doc, "ARTICLE 7 \u2014 CONFIDENTIALIT\u00c9")
doc.add_paragraph().add_run("Chaque partie garde confidentiels les documents de l\u2019autre partie.").font.name = "Arial"
titre(doc, "ARTICLE 8 \u2014 RGPD")
p = doc.add_paragraph(); champ(p, "NomClient"); txt(p, " est Responsable de traitement. SGI est Sous-traitant (art. 28 RGPD).")
titre(doc, "ARTICLE 9 \u2014 DROIT APPLICABLE")
doc.add_paragraph().add_run("Tribunal de Commerce de Lille comp\u00e9tent.").font.name = "Arial"
bloc_signataires(doc)
doc.add_page_break()
p = doc.add_paragraph(); txt(p, "ANNEXE 1 \u2014 D\u00e9tail des progiciels", bold=True, color=BLUE, size=13)
p = doc.add_paragraph(); txt(p, "au contrat "); champ(p, "LibCodePrestation"); txt(p, " "); champ(p, "NoContrat")
t3 = doc.add_table(rows=4, cols=3); t3.style = "Table Grid"
entete_tableau(t3, "Progiciel", "Description", "Prix HT/an")
for i, n in enumerate([1,2], 1):
    for j, cn in enumerate([f"Progiciel{n}", f"DescriptionProgiciel{n}", f"PrixProgiciel{n}"], 0):
        p2 = t3.rows[i].cells[j].paragraphs[0]; p2.clear(); champ(p2, cn)
p0 = t3.rows[3].cells[0].paragraphs[0]; p0.clear(); txt(p0, "TOTAL HT", bold=True)
t3.rows[3].cells[1].paragraphs[0].clear()
p1 = t3.rows[3].cells[2].paragraphs[0]; p1.clear(); champ(p1, "TotalHT"); txt(p1, " \u20ac")
doc.add_page_break()
p = doc.add_paragraph(); txt(p, "ANNEXE 2 \u2014 Collectivit\u00e9s rattach\u00e9es", bold=True, color=BLUE, size=13)
p = doc.add_paragraph(); txt(p, "au contrat "); champ(p, "LibCodePrestation"); txt(p, " "); champ(p, "NoContrat")
t4 = doc.add_table(rows=11, cols=2); t4.style = "Table Grid"
entete_tableau(t4, "R\u00e9f. site", "Nom de la collectivit\u00e9")
for i in range(1, 11):
    p0 = t4.rows[i].cells[0].paragraphs[0]; p0.clear(); champ(p0, f"COL{i}IdSite")
    p1 = t4.rows[i].cells[1].paragraphs[0]; p1.clear(); champ(p1, f"COL{i}NomSite")
doc.add_page_break()
p = doc.add_paragraph(); txt(p, "ANNEXE 3 \u2014 RGPD", bold=True, color=BLUE, size=13)
p = doc.add_paragraph(); txt(p, "au contrat "); champ(p, "LibCodePrestation"); txt(p, " "); champ(p, "NoContrat")
titre(doc, "1. R\u00f4les")
p = doc.add_paragraph(); champ(p, "NomClient"); txt(p, " : Responsable de traitement. SGI : Sous-traitant (art. 28 RGPD).")
titre(doc, "2. Finalit\u00e9s")
doc.add_paragraph().add_run("Traitement des donn\u00e9es uniquement pour l\u2019ex\u00e9cution du contrat.").font.name = "Arial"
titre(doc, "3. S\u00e9curit\u00e9")
doc.add_paragraph().add_run("Mesures techniques et organisationnelles appropri\u00e9es mises en oeuvre par SGI.").font.name = "Arial"
p = doc.add_paragraph(); txt(p, "Fait \u00e0 Villeneuve d\u2019Ascq, le "); champ(p, "DateSignature")
doc.save(str(OUT / "Modele_Contrat_Cosoluce_et_Annexes.docx"))
print("OK Cosoluce")

# ═══════════════════════════════════════════════════
# CANTINE
# ═══════════════════════════════════════════════════
doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2); s.bottom_margin=Cm(2); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
p = doc.add_paragraph(); txt(p, "SGI INFORMATIQUE \u2014 Contrat Cantine de France", bold=True, color=BLUE, size=13)
doc.add_paragraph()
entete_client(doc, [("R\u00e9f. client : ", "RefClient")])
doc.add_paragraph()
p = doc.add_paragraph(); txt(p, "Contrat N\u00b0 : "); champ(p, "NoContrat"); txt(p, " \u2014 Date : "); champ(p, "DateDoc")
titre(doc, "ARTICLE 1 \u2014 OBJET")
doc.add_paragraph().add_run("Abonnement au service Cantine de France pour la gestion de la restauration scolaire.").font.name = "Arial"
titre(doc, "ARTICLE 2 \u2014 DUR\u00c9E")
p = doc.add_paragraph(); txt(p, "Du "); champ(p, "DateDebut"); txt(p, " au "); champ(p, "DateFin"); txt(p, ". Reconduction annuelle tacite.")
titre(doc, "ARTICLE 3 \u2014 PRIX")
t2 = doc.add_table(rows=2, cols=2); t2.style = "Table Grid"
entete_tableau(t2, "Prestation", "Montant HT/an")
p0 = t2.rows[1].cells[0].paragraphs[0]; p0.clear(); txt(p0, "Abonnement Cantine de France")
p1 = t2.rows[1].cells[1].paragraphs[0]; p1.clear(); champ(p1, "MontantHT"); txt(p1, " \u20ac")
p = doc.add_paragraph(); txt(p, "Premi\u00e8re facturation le "); champ(p, "DatePremFactu"); txt(p, ".")
titre(doc, "ARTICLE 4 \u2014 R\u00c9VISION")
doc.add_paragraph().add_run("R\u00e9vision annuelle Syntec Octobre : Prix N = Prix N-1 \u00d7 (Oct N-1 / Oct N-2).").font.name = "Arial"
titre(doc, "ARTICLE 5 \u2014 R\u00c9SILIATION")
doc.add_paragraph().add_run("Pr\u00e9avis 3 mois par lettre recommand\u00e9e avant \u00e9ch\u00e9ance.").font.name = "Arial"
titre(doc, "ARTICLE 6 \u2014 RGPD")
p = doc.add_paragraph(); champ(p, "NomClient"); txt(p, " est Responsable de traitement. SGI est Sous-traitant.")
titre(doc, "ARTICLE 7 \u2014 DROIT APPLICABLE")
doc.add_paragraph().add_run("Tribunal de Commerce de Lille.").font.name = "Arial"
bloc_signataires(doc)
doc.save(str(OUT / "Modele_Contrat_Cantine_de_France.docx"))
print("OK Cantine")

# ═══════════════════════════════════════════════════
# MAINTENANCE
# ═══════════════════════════════════════════════════
doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2); s.bottom_margin=Cm(2); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
p = doc.add_paragraph(); txt(p, "SGI INFORMATIQUE \u2014 Contrat Maintenance Syst\u00e8me", bold=True, color=BLUE, size=13)
doc.add_paragraph()
entete_client(doc, [("T\u00e9l : ", "TelClient")])
doc.add_paragraph()
p = doc.add_paragraph(); txt(p, "Contrat N\u00b0 : "); champ(p, "NoContrat"); txt(p, " \u2014 Date : "); champ(p, "DateDoc")
titre(doc, "ARTICLE 1 \u2014 OBJET")
p = doc.add_paragraph(); txt(p, "Maintenance du syst\u00e8me informatique de "); champ(p, "NomClient"); txt(p, ".")
titre(doc, "ARTICLE 2 \u2014 DUR\u00c9E")
p = doc.add_paragraph(); txt(p, "Du "); champ(p, "DateDebut"); txt(p, " au "); champ(p, "DateFin"); txt(p, ". Reconduction expresse par p\u00e9riodes d\u2019une ou plusieurs ann\u00e9es.")
titre(doc, "ARTICLE 3 \u2014 PRESTATIONS")
doc.add_paragraph().add_run("Maintenance pr\u00e9ventive et corrective, assistance t\u00e9l\u00e9phonique, intervention sur site, MAJ syst\u00e8mes.").font.name = "Arial"
titre(doc, "ARTICLE 4 \u2014 CONDITIONS FINANCI\u00c8RES")
t2 = doc.add_table(rows=2, cols=2); t2.style = "Table Grid"
entete_tableau(t2, "Prestation", "Montant HT/an")
p0 = t2.rows[1].cells[0].paragraphs[0]; p0.clear(); txt(p0, "Maintenance syst\u00e8me")
p1 = t2.rows[1].cells[1].paragraphs[0]; p1.clear(); champ(p1, "MontantHT"); txt(p1, " \u20ac")
titre(doc, "ARTICLE 5 \u2014 R\u00c9VISION")
doc.add_paragraph().add_run("R\u00e9vision annuelle Syntec Ao\u00fbt.").font.name = "Arial"
titre(doc, "ARTICLE 6 \u2014 EXCLUSIVIT\u00c9")
doc.add_paragraph().add_run("Toute prestation de maintenance doit \u00eatre exclusivement confi\u00e9e au Prestataire.").font.name = "Arial"
titre(doc, "ARTICLE 7 \u2014 R\u00c9SILIATION")
doc.add_paragraph().add_run("Pr\u00e9avis 3 mois par lettre recommand\u00e9e avant \u00e9ch\u00e9ance.").font.name = "Arial"
titre(doc, "ARTICLE 8 \u2014 DROIT APPLICABLE")
doc.add_paragraph().add_run("Tribunal du si\u00e8ge social du Prestataire comp\u00e9tent.").font.name = "Arial"
bloc_signataires(doc)
doc.save(str(OUT / "Modele_Contrat_Maintenance_Systeme.docx"))
print("OK Maintenance")

# ═══════════════════════════════════════════════════
# ASSISTANCE CITYWEB
# ═══════════════════════════════════════════════════
doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2); s.bottom_margin=Cm(2); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
p = doc.add_paragraph(); txt(p, "SGI INFORMATIQUE \u2014 Contrat Assistance T\u00e9l\u00e9phonique Cityweb", bold=True, color=BLUE, size=13)
doc.add_paragraph()
entete_client(doc, [("Interlocuteur : ", "NomInterlocuteur"), ("T\u00e9l : ", "TelClient")])
doc.add_paragraph()
p = doc.add_paragraph(); txt(p, "Contrat N\u00b0 : "); champ(p, "NoContrat"); txt(p, " \u2014 N\u00b0 abonnement : "); champ(p, "NoAbonnement"); txt(p, " \u2014 Date : "); champ(p, "DateDoc")
titre(doc, "ARTICLE 1 \u2014 OBJET")
doc.add_paragraph().add_run("Assistance t\u00e9l\u00e9phonique sur les logiciels Cityweb pour les agents de la collectivit\u00e9.").font.name = "Arial"
titre(doc, "ARTICLE 2 \u2014 DUR\u00c9E")
p = doc.add_paragraph(); txt(p, "Du "); champ(p, "DateDebut"); txt(p, " au "); champ(p, "DateFin"); txt(p, ". Reconduction annuelle tacite.")
titre(doc, "ARTICLE 3 \u2014 PRIX")
t2 = doc.add_table(rows=2, cols=2); t2.style = "Table Grid"
entete_tableau(t2, "Prestation", "Montant HT/an")
p0 = t2.rows[1].cells[0].paragraphs[0]; p0.clear(); txt(p0, "Assistance t\u00e9l\u00e9phonique Cityweb")
p1 = t2.rows[1].cells[1].paragraphs[0]; p1.clear(); champ(p1, "MontantHT"); txt(p1, " \u20ac")
titre(doc, "ARTICLE 4 \u2014 MODALIT\u00c9S")
doc.add_paragraph().add_run("Assistance aux heures ouvrables, acc\u00e8s t\u00e9l\u00e9phone et ticketing.").font.name = "Arial"
titre(doc, "ARTICLE 5 \u2014 R\u00c9VISION")
doc.add_paragraph().add_run("R\u00e9vision annuelle Syntec Ao\u00fbt.").font.name = "Arial"
titre(doc, "ARTICLE 6 \u2014 R\u00c9SILIATION")
doc.add_paragraph().add_run("Pr\u00e9avis 3 mois par lettre recommand\u00e9e.").font.name = "Arial"
titre(doc, "ARTICLE 7 \u2014 DROIT APPLICABLE")
doc.add_paragraph().add_run("Tribunal de Commerce de Lille.").font.name = "Arial"
bloc_signataires(doc)
doc.save(str(OUT / "Modele_Contrat_Assistance_Cityweb.docx"))
print("OK Cityweb")

# ═══════════════════════════════════════════════════
# ASSISTANCE COSOLUCE
# ═══════════════════════════════════════════════════
doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2); s.bottom_margin=Cm(2); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
p = doc.add_paragraph(); txt(p, "SGI INFORMATIQUE \u2014 Contrat Assistance T\u00e9l\u00e9phonique Cosoluce", bold=True, color=BLUE, size=13)
doc.add_paragraph()
entete_client(doc, [("N\u00b0 contrat Cosoluce : ", "NoContratCosoluce")])
doc.add_paragraph()
p = doc.add_paragraph(); txt(p, "Contrat N\u00b0 : "); champ(p, "NoContrat"); txt(p, " \u2014 Date : "); champ(p, "DateDoc")
titre(doc, "ARTICLE 1 \u2014 OBJET")
doc.add_paragraph().add_run("Assistance t\u00e9l\u00e9phonique sur les progiciels Cosoluce.").font.name = "Arial"
titre(doc, "ARTICLE 2 \u2014 DUR\u00c9E")
p = doc.add_paragraph(); txt(p, "Du "); champ(p, "DateDebut"); txt(p, " au "); champ(p, "DateFin"); txt(p, ". Reconduction annuelle tacite.")
titre(doc, "ARTICLE 3 \u2014 PRIX")
t2 = doc.add_table(rows=2, cols=2); t2.style = "Table Grid"
entete_tableau(t2, "Prestation", "Montant HT/an")
p0 = t2.rows[1].cells[0].paragraphs[0]; p0.clear(); txt(p0, "Assistance t\u00e9l\u00e9phonique Cosoluce")
p1 = t2.rows[1].cells[1].paragraphs[0]; p1.clear(); champ(p1, "MontantHT"); txt(p1, " \u20ac")
titre(doc, "ARTICLE 4 \u2014 R\u00c9VISION")
doc.add_paragraph().add_run("R\u00e9vision annuelle Syntec Ao\u00fbt.").font.name = "Arial"
titre(doc, "ARTICLE 5 \u2014 R\u00c9SILIATION")
doc.add_paragraph().add_run("Pr\u00e9avis 3 mois par lettre recommand\u00e9e.").font.name = "Arial"
titre(doc, "ARTICLE 6 \u2014 DROIT APPLICABLE")
doc.add_paragraph().add_run("Tribunal de Commerce de Lille.").font.name = "Arial"
bloc_signataires(doc)
doc.save(str(OUT / "Modele_Contrat_Assistance_Cosoluce.docx"))
print("OK Cosoluce Assistance")

print("\nTous les modeles generes dans", str(OUT))
